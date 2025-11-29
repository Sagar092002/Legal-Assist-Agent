"""
Template Converter Service
Automatically converts user-uploaded templates to Jinja2 format

Features:
- Detects various placeholder types (hash #, underscores _, dots ., brackets [], etc.)
- Converts to Jinja2 format {{ variable_name }}
- Generates intelligent variable names using AI
- Extracts field metadata
- Validates conversion quality
"""

import re
import json
import logging
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from .azure_openai_service import ai_service
from .config import AIConfig

logger = logging.getLogger(__name__)


class TemplateConverter:
    """Converts user templates to Jinja2 format"""
    
    # Placeholder detection patterns
    PLACEHOLDER_PATTERNS = {
        'hash': r'#\d+',                          # #1, #2, #123
        'underscore': r'_{4,}',                   # ____, ___________
        'dots': r'\.{4,}',                        # ...., ..........
        'brackets_square': r'\[[\s\w-]+\]',       # [NAME], [DATE]
        'brackets_curly': r'\{[\s\w-]+\}',        # {NAME}, {DATE}
        'brackets_angle': r'<[\s\w-]+>',          # <NAME>, <DATE>
        'dollar': r'\$\{[\w_]+\}',                # ${NAME}, ${DATE}
        'percent': r'%[\w_]+%',                   # %NAME%, %DATE%
    }
    
    def __init__(self):
        self.ai_enabled = AIConfig.validate()
        if not self.ai_enabled:
            logger.warning("AI service not configured - using basic variable naming")
    
    def analyze_template(self, doc_path: str) -> Dict:
        """
        Analyze template to detect placeholders
        
        Returns:
            {
                'total_placeholders': 15,
                'placeholder_types': {
                    'hash': ['#1', '#2', ...],
                    'underscore': ['____', '________'],
                    ...
                },
                'suggested_conversions': {
                    '#1': 'party_name_1',
                    '#2': 'effective_date',
                    ...
                },
                'context': {
                    '#1': 'appears near "Party Name:"',
                    ...
                }
            }
        """
        try:
            doc = Document(doc_path)
            
            # Extract all text for analysis
            all_text = []
            placeholder_contexts = {}
            
            # Analyze paragraphs
            for para in doc.paragraphs:
                all_text.append(para.text)
            
            # Analyze tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            all_text.append(para.text)
            
            full_text = '\n'.join(all_text)
            
            # Detect placeholders by type
            detected_placeholders = {}
            for ptype, pattern in self.PLACEHOLDER_PATTERNS.items():
                matches = re.findall(pattern, full_text)
                if matches:
                    # Remove duplicates while preserving order
                    detected_placeholders[ptype] = list(dict.fromkeys(matches))
            
            # Get context for each placeholder
            for ptype, placeholders in detected_placeholders.items():
                for placeholder in placeholders:
                    context = self._get_placeholder_context(full_text, placeholder)
                    placeholder_contexts[placeholder] = context
            
            # Generate suggested variable names using AI
            suggested_conversions = {}
            if self.ai_enabled and placeholder_contexts:
                suggested_conversions = self._generate_variable_names_ai(
                    placeholder_contexts,
                    full_text[:2000]  # Send first 2000 chars for context
                )
            else:
                # Fallback to basic naming
                suggested_conversions = self._generate_variable_names_basic(detected_placeholders)
            
            total = sum(len(v) for v in detected_placeholders.values())
            
            return {
                'success': True,
                'total_placeholders': total,
                'placeholder_types': detected_placeholders,
                'suggested_conversions': suggested_conversions,
                'context': placeholder_contexts,
                'document_preview': full_text[:500]
            }
            
        except Exception as e:
            logger.error(f"Template analysis failed: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def convert_to_jinja2(
        self,
        doc_path: str,
        output_path: str,
        variable_mapping: Dict[str, str]
    ) -> Dict:
        """
        Convert template placeholders to Jinja2 format
        
        Args:
            doc_path: Path to original template
            output_path: Path to save converted template
            variable_mapping: Dict mapping placeholders to variable names
                Example: {'#1': 'party_name_1', '#2': 'effective_date'}
        
        Returns:
            {
                'success': True,
                'converted_count': 15,
                'output_path': '...',
                'remaining_placeholders': []
            }
        """
        try:
            doc = Document(doc_path)
            converted_count = 0
            
            # Sort mappings by length (descending) to avoid partial replacements
            sorted_mappings = sorted(
                variable_mapping.items(),
                key=lambda x: len(x[0]),
                reverse=True
            )
            
            # Convert paragraphs
            for para in doc.paragraphs:
                original_text = para.text
                for placeholder, var_name in sorted_mappings:
                    if placeholder in para.text:
                        # Escape special regex characters in placeholder
                        escaped = re.escape(placeholder)
                        # Replace with Jinja2 syntax
                        para.text = re.sub(
                            escaped,
                            f"{{{{ {var_name} }}}}",
                            para.text
                        )
                        if para.text != original_text:
                            converted_count += 1
                            logger.info(f"Converted {placeholder} → {{{{ {var_name} }}}}")
            
            # Convert tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            original_text = para.text
                            for placeholder, var_name in sorted_mappings:
                                if placeholder in para.text:
                                    escaped = re.escape(placeholder)
                                    para.text = re.sub(
                                        escaped,
                                        f"{{{{ {var_name} }}}}",
                                        para.text
                                    )
                                    if para.text != original_text:
                                        converted_count += 1
                                        logger.info(f"Converted {placeholder} → {{{{ {var_name} }}}} (table)")
            
            # Save converted template
            doc.save(output_path)
            
            # Check for remaining unconverted placeholders
            remaining = self._find_remaining_placeholders(output_path)
            
            return {
                'success': True,
                'converted_count': converted_count,
                'output_path': output_path,
                'remaining_placeholders': remaining,
                'message': f'Successfully converted {converted_count} placeholders'
            }
            
        except Exception as e:
            logger.error(f"Template conversion failed: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _get_placeholder_context(self, text: str, placeholder: str, context_chars: int = 50) -> str:
        """Get surrounding text context for a placeholder"""
        try:
            index = text.find(placeholder)
            if index == -1:
                return ""
            
            start = max(0, index - context_chars)
            end = min(len(text), index + len(placeholder) + context_chars)
            
            context = text[start:end]
            # Clean up
            context = ' '.join(context.split())
            return context
        except:
            return ""
    
    def _generate_variable_names_ai(
        self,
        placeholder_contexts: Dict[str, str],
        document_preview: str
    ) -> Dict[str, str]:
        """Use AI to generate meaningful variable names based on context"""
        try:
            # Prepare contexts for GPT
            contexts_list = [
                f"'{placeholder}' appears in: \"{context}\""
                for placeholder, context in list(placeholder_contexts.items())[:20]  # Limit to 20
            ]
            
            prompt = f"""You are analyzing a legal document template. Generate appropriate variable names for placeholders.

Document preview:
{document_preview}

Placeholders with context:
{chr(10).join(contexts_list)}

Generate clean, descriptive variable names (snake_case, no special chars).

Rules:
1. Use context to infer meaning
2. Common legal fields: party_name, effective_date, amount, address, etc.
3. For numbered placeholders (#1, #2), use descriptive names
4. Keep names under 30 characters
5. Use snake_case (lowercase with underscores)

Return ONLY a JSON object mapping placeholders to variable names:
{{
    "#1": "party_name_1",
    "#2": "effective_date",
    "____": "recipient_address",
    ...
}}"""

            response = ai_service.chat_completion([
                {"role": "system", "content": "You are a legal document template analyzer. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ], temperature=0.3, max_tokens=1000)
            
            # Parse JSON response
            match = re.search(r'\{.*\}', response, re.DOTALL)
            if match:
                mapping = json.loads(match.group(0))
                logger.info(f"✅ AI generated {len(mapping)} variable names")
                return mapping
            else:
                logger.warning("AI response not in JSON format, using fallback")
                return self._generate_variable_names_basic({'mixed': list(placeholder_contexts.keys())})
                
        except Exception as e:
            logger.error(f"AI variable name generation failed: {e}")
            return self._generate_variable_names_basic({'mixed': list(placeholder_contexts.keys())})
    
    def _generate_variable_names_basic(self, detected_placeholders: Dict) -> Dict[str, str]:
        """Fallback: Generate basic variable names without AI"""
        mapping = {}
        counter = 1
        
        for ptype, placeholders in detected_placeholders.items():
            for placeholder in placeholders:
                # Try to extract meaningful name from bracket placeholders
                if ptype in ['brackets_square', 'brackets_curly', 'brackets_angle']:
                    # Extract text inside brackets
                    inner = re.sub(r'[\[\]<>{}]', '', placeholder).strip()
                    var_name = inner.lower().replace(' ', '_').replace('-', '_')
                    mapping[placeholder] = var_name
                else:
                    # Generic naming
                    mapping[placeholder] = f"field_{counter}"
                    counter += 1
        
        return mapping
    
    def _find_remaining_placeholders(self, doc_path: str) -> List[str]:
        """Check for any remaining unconverted placeholders"""
        try:
            doc = Document(doc_path)
            all_text = []
            
            for para in doc.paragraphs:
                all_text.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            all_text.append(para.text)
            
            full_text = '\n'.join(all_text)
            
            # Check for remaining placeholders
            remaining = []
            for ptype, pattern in self.PLACEHOLDER_PATTERNS.items():
                matches = re.findall(pattern, full_text)
                remaining.extend(matches)
            
            return list(dict.fromkeys(remaining))  # Remove duplicates
            
        except:
            return []
    
    def validate_conversion(self, doc_path: str) -> Dict:
        """Validate that conversion was successful"""
        try:
            doc = Document(doc_path)
            
            # Count Jinja2 variables
            all_text = []
            for para in doc.paragraphs:
                all_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            all_text.append(para.text)
            
            full_text = '\n'.join(all_text)
            
            # Find all Jinja2 variables
            jinja_vars = re.findall(r'\{\{\s*(\w+)\s*\}\}', full_text)
            unique_vars = list(dict.fromkeys(jinja_vars))
            
            # Check for remaining placeholders
            remaining = self._find_remaining_placeholders(doc_path)
            
            return {
                'success': True,
                'jinja_variable_count': len(jinja_vars),
                'unique_variables': unique_vars,
                'remaining_placeholders': remaining,
                'is_valid': len(remaining) == 0 and len(unique_vars) > 0
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }
    
    def extract_template_metadata(
        self,
        doc_path: str,
        template_name: str,
        category: str = "Custom"
    ) -> Dict:
        """Extract metadata for template configuration"""
        try:
            validation = self.validate_conversion(doc_path)
            
            if not validation['success']:
                return validation
            
            # Build field configuration
            fields = {}
            for var_name in validation['unique_variables']:
                # Convert snake_case to title case for label
                label = var_name.replace('_', ' ').title()
                
                # Detect field type based on name
                field_type = 'text'
                if any(word in var_name.lower() for word in ['date', 'day', 'month', 'year']):
                    field_type = 'date'
                elif any(word in var_name.lower() for word in ['amount', 'price', 'rent', 'fee']):
                    field_type = 'number'
                elif 'email' in var_name.lower():
                    field_type = 'email'
                
                # Mark as required by default
                fields[var_name] = {
                    'label': label,
                    'type': field_type,
                    'required': True,
                    'example': ''
                }
            
            metadata = {
                'name': template_name,
                'category': category,
                'filename': Path(doc_path).name,
                'description': f'User-uploaded template: {template_name}',
                'keywords': [template_name.lower(), category.lower()],
                'fields': fields,
                'is_user_template': True,
                'jinja_variable_count': validation['jinja_variable_count']
            }
            
            return {
                'success': True,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Metadata extraction failed: {e}")
            return {
                'success': False,
                'error': str(e)
            }


# Singleton instance
template_converter = TemplateConverter()
