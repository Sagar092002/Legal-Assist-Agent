"""
Convert Lease-Deed template to Jinja2 format
"""
from docx import Document
from docxtpl import DocxTemplate
import os

def create_jinja2_lease_template():
    """Create a new Lease Agreement template with Jinja2 placeholders"""
    
    # Create a new document
    doc = Document()
    
    # Title
    doc.add_heading('DEED OF LEASE (RENT AGREEMENT)', 0)
    
    # Opening paragraph
    opening = doc.add_paragraph()
    opening.add_run(
        'This Deed of Lease is made at {{ location }} this {{ day }} day of {{ month }}, {{ year }} '
        'between {{ lessor_name }} of {{ lessor_address }} hereinafter called "The Lessor" of the One Part '
        'and {{ lessee_name }} of {{ lessee_address }} hereinafter called "The Lessee" of the Other Part.'
    )
    
    # WHEREAS clauses
    doc.add_paragraph(
        'WHEREAS, the Lessor is absolutely seized and possessed of or otherwise well and sufficiently '
        'entitled to the land and premises described in the Schedule hereunder written.'
    )
    
    doc.add_paragraph(
        'AND WHEREAS, the Lessor has agreed to grant to the Lessee a lease in respect of the said land '
        'and premises for a term of {{ lease_duration_years }} years in the manner hereinafter appearing.'
    )
    
    doc.add_heading('NOW THIS DEED WITNESSETH AS FOLLOWS:', level=2)
    
    # Clause 1 - Demise and Rent
    clause1 = doc.add_paragraph()
    clause1.add_run('1. ').bold = True
    clause1.add_run(
        'In pursuance of the said agreement and in consideration of the rent hereby reserved and of the '
        'terms and conditions, covenants and agreements herein contained and on the part of the Lessee '
        'to be observed and performed the Lessor doth hereby demise unto the Lessee all that the said land '
        'and premises situated at {{ property_address }} and described in the Schedule hereunder written '
        '(hereinafter for the brevity\'s sake referred to as "the demised premises") to hold the demised '
        'premises unto the Lessee (and his heirs, executors, administrators and assigns) for a term of '
        '{{ lease_duration_years }} years commencing from {{ lease_start_date }}, but subject to the earlier '
        'determination of this demise as hereinafter provided and yielding and paying therefor during the said '
        'term the monthly ground rent of ₹{{ monthly_rent }} free and clear of all deductions and strictly in '
        'advance on or before the {{ payment_day }}th day of each and every calendar month. The first of such '
        'monthly ground rent shall be paid on {{ first_payment_date }} and the subsequent rent to be paid on or '
        'before the {{ payment_day }}th day of every succeeding month regularly.'
    )
    
    # Clause 2 - Lessee's Covenants
    clause2 = doc.add_paragraph()
    clause2.add_run('2. ').bold = True
    clause2.add_run(
        'The Lessee hereby for himself, his heirs, executors, administrators and assigns and to the intent '
        'that the obligations herein contained shall continue throughout the term hereby created covenants '
        'with the Lessor as follows:'
    )
    
    # Sub-clauses
    doc.add_paragraph(
        'a. To pay the ground rent hereby reserved on the days and in the manner aforesaid clear of all '
        'deductions. The first of such monthly rent as hereinbefore provided shall be paid on '
        '{{ first_payment_date }} and the subsequent rent shall be paid on the {{ payment_day }}th day of '
        'every succeeding month regularly and if the ground rent is not paid on the due dates the Lessee '
        'shall pay interest thereon at the rate of {{ late_fee_rate }}% per annum from the due date till '
        'payment, though the payment of interest shall not entitle the Lessee to make default in payment '
        'of rent on due dates.'
    )
    
    doc.add_paragraph(
        'b. To bear pay and discharge the existing and future rates, taxes and assessment duties, cess, '
        'impositions, outgoing and burdens whatsoever which may at any time or from time to time during '
        'the term hereby created be imposed or charged upon the demised land and the building or structures '
        'standing thereon and on the buildings or structures hereafter to be erected and for the time being '
        'standing on the demised land and payable either by the owners, occupiers or tenants thereof and to '
        'keep the Lessor and his estate and effects indemnified against all such payment.'
    )
    
    doc.add_paragraph(
        'c. To keep the buildings and structures on the demised premises in good and tenantable repairs in '
        'the same way as the Lessor is liable to do under the law provided that if the Lessee so desires he '
        'shall have power to demolish any existing building or structure without being accountable to the '
        'Lessor for the building material of such building and structure and the Lessee shall have also power '
        'to construct any new buildings in their place.'
    )
    
    doc.add_paragraph(
        'd. The Lessee shall be at liberty to carry out any additions or alterations to the buildings or '
        'structures at present existing on the demised premises or to put up any additional structures or '
        'buildings on the demised premises in accordance with the plans approved by the authorities at any '
        'time or from time to time during the subsistence of the term hereby created.'
    )
    
    doc.add_paragraph(
        'e. Not to sell or dispose of any earth, gravel or sand from the demised land and not to excavate '
        'the same except so far as may be necessary for the execution of construction work.'
    )
    
    doc.add_paragraph(
        'f. To use or permit to be used the buildings and structures to be constructed on the demised '
        'premises for any and all lawful purposes as may be permitted by the authorities from time to time.'
    )
    
    # Clause 3 - Lessor's Covenants
    clause3 = doc.add_paragraph()
    clause3.add_run('3. ').bold = True
    clause3.add_run('The Lessor doth hereby covenant with the Lessee that:')
    
    doc.add_paragraph(
        'a. The Lessor now has in himself good right full power and absolute authority to demise unto the '
        'Lessee the demised premises and the buildings and structures standing thereon in the manner herein '
        'appearing.'
    )
    
    doc.add_paragraph(
        'b. That on the Lessee paying the said monthly ground rent on the due dates thereof and in the manner '
        'herein provided and observing and performing the covenants, conditions, and stipulations herein '
        'contained and on his part to be observed and performed shall and may peaceably and quietly hold, '
        'possess and enjoy the demised premises together with the buildings and structures standing thereon '
        'during the term hereby created without any eviction, interruption, disturbance, claim and demand '
        'whatsoever by the Lessor or any person or persons lawfully or equitably claiming by, from, under or '
        'in trust for him.'
    )
    
    # Clause 4 - Termination
    clause4 = doc.add_paragraph()
    clause4.add_run('4. ').bold = True
    clause4.add_run(
        'It is hereby agreed and declared that these presents are granted on the express condition that if '
        'the said monthly ground rent or any part thereof payable in the manner hereinbefore mentioned shall '
        'be in arrears for the space of {{ arrears_months }} months after the same shall have become due and '
        'payable then and in such event it shall be lawful for the Lessor to terminate this lease with '
        '{{ notice_period }} months notice in writing to the Lessee.'
    )
    
    # Clause 5 - Final provisions
    clause5 = doc.add_paragraph()
    clause5.add_run('5. ').bold = True
    clause5.add_run('And it is hereby expressly agreed and declared between the parties as follows:')
    
    doc.add_paragraph(
        'a. On the expiration of the term hereby created or earlier determination under the provisions '
        'hereof all the buildings and structures standing on the demised land shall automatically vest in '
        'the Lessor without payment of any compensation therefor by the Lessor to the Lessee.'
    )
    
    doc.add_paragraph(
        'b. The Lessee shall not be entitled, without obtaining in writing the permission of the Lessor, '
        'to assign mortgage, sublet (except to the extent of creating monthly tenancies) or otherwise part '
        'with possession of the demised premises or any of them or any part thereof and the buildings and '
        'structure standing thereon though such permission shall not be unreasonably withheld.'
    )
    
    # Witness clause
    doc.add_paragraph()
    witness = doc.add_paragraph()
    witness.add_run('IN WITNESS WHEREOF ').bold = True
    witness.add_run(
        'the Lessor and the Lessee have put their respective hands on the original and duplicate hereof '
        'the day and year first herein above written.'
    )
    
    # Schedule
    doc.add_heading('THE SCHEDULE ABOVE REFERRED TO', level=2)
    doc.add_paragraph('{{ property_description }}')
    
    # Signatures
    doc.add_paragraph('\n\n')
    doc.add_paragraph('Signed and delivered by the Lessor')
    doc.add_paragraph('{{ lessor_name }}')
    doc.add_paragraph('\nIn the presence of:')
    doc.add_paragraph('Witness 1: {{ lessor_witness_1 }}')
    doc.add_paragraph('Witness 2: {{ lessor_witness_2 }}')
    
    doc.add_paragraph('\n\n')
    doc.add_paragraph('Signed and delivered by the Lessee')
    doc.add_paragraph('{{ lessee_name }}')
    doc.add_paragraph('\nIn the presence of:')
    doc.add_paragraph('Witness 1: {{ lessee_witness_1 }}')
    doc.add_paragraph('Witness 2: {{ lessee_witness_2 }}')
    
    # Save as docx template
    output_path = os.path.join(os.path.dirname(__file__), '..', 'data', 'templates', 
                               'Lease-Agreement-Jinja2.docx')
    doc.save(output_path)
    
    print(f"✅ Created Jinja2 template: {output_path}")
    
    # List all variables used
    variables = [
        'location', 'day', 'month', 'year',
        'lessor_name', 'lessor_address',
        'lessee_name', 'lessee_address',
        'lease_duration_years',
        'property_address',
        'lease_start_date',
        'monthly_rent',
        'payment_day',
        'first_payment_date',
        'late_fee_rate',
        'arrears_months',
        'notice_period',
        'property_description',
        'lessor_witness_1', 'lessor_witness_2',
        'lessee_witness_1', 'lessee_witness_2'
    ]
    
    print(f"\n📋 Template Variables ({len(variables)}):")
    for var in variables:
        print(f"   • {{ {var} }}")
    
    return variables

if __name__ == '__main__':
    create_jinja2_lease_template()
