"""
Fix NDA-Jinja2.docx - complete the conversion
"""

from docx import Document
from pathlib import Path

BASE_DIR = Path(__file__).parent.parent
TEMPLATES_DIR = BASE_DIR / "data" / "templates"
input_file = TEMPLATES_DIR / "NDA-Jinja2.docx"
output_file = TEMPLATES_DIR / "NDA-Jinja2-Fixed.docx"

print("="*80)
print("FIXING NDA-JINJA2 TEMPLATE")
print("="*80)

doc = Document(input_file)

# Additional replacements needed
replacements = {
    '"____"': '"{{ party_1_short_name }}"',
    '[Please fill in Customer name]': '{{ party_2_name }}',
    '[Please fill in Customers name]': '{{ party_2_name }}',
    'Désignations': 'Designation',
}

conversions = []

# Process paragraphs
for para in doc.paragraphs:
    original = para.text
    new_text = original
    
    for old, new in replacements.items():
        if old in new_text:
            new_text = new_text.replace(old, new)
            conversions.append(f"{old} → {new}")
    
    if new_text != original:
        para.text = new_text

# Process tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                original = para.text
                new_text = original
                
                for old, new in replacements.items():
                    if old in new_text:
                        new_text = new_text.replace(old, new)
                        conversions.append(f"{old} → {new} (table)")
                
                if new_text != original:
                    para.text = new_text

# Save
doc.save(output_file)

print("\n✅ Additional conversions:")
for conv in set(conversions):
    print(f"  ✓ {conv}")

print(f"\n✅ Fixed and saved to: {output_file}")

# Verify all Jinja2 variables
import re
text = '\n'.join([p.text for p in doc.paragraphs])
jinja_vars = set(re.findall(r'\{\{\s*(\w+)\s*\}\}', text))

print("\n📋 All Jinja2 variables in template:")
for var in sorted(jinja_vars):
    print(f"  - {{ {var} }}")

print(f"\nTotal variables: {len(jinja_vars)}")

# Check for any remaining placeholders
remaining_placeholders = []
if '<<' in text or '>>' in text:
    remaining_placeholders.append("Found << or >>")
if '[Please fill in' in text:
    remaining_placeholders.append("Found [Please fill in ...]")
if '"____"' in text:
    remaining_placeholders.append('Found "____"')

if remaining_placeholders:
    print("\n⚠️ WARNING: Remaining unconverted placeholders:")
    for p in remaining_placeholders:
        print(f"  ! {p}")
else:
    print("\n✅ No remaining placeholders - conversion complete!")
