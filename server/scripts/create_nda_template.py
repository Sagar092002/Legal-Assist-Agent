"""
Create a professional NDA Jinja2 template
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

BASE_DIR = Path(__file__).parent.parent
TEMPLATES_DIR = BASE_DIR / "data" / "templates" / "employment"
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

# Create new document
doc = Document()

# Title
title = doc.add_paragraph()
title_run = title.add_run("NON-DISCLOSURE AGREEMENT")
title_run.bold = True
title_run.font.size = Pt(16)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Agreement details
doc.add_paragraph(f"This Non-Disclosure Agreement (\"Agreement\") is entered into on {{{{ effective_date }}}} (\"Effective Date\")")

doc.add_paragraph(f"BETWEEN")

doc.add_paragraph(f"""{{{{ disclosing_party }}}}, having its principal place of business at {{{{ disclosing_party_address }}}} (hereinafter referred to as the \"Disclosing Party\", which expression shall, unless repugnant to the context or meaning thereof, be deemed to include its successors and permitted assigns)""")

doc.add_paragraph("AND")

doc.add_paragraph(f"""{{{{ receiving_party }}}}, having its principal place of business at {{{{ receiving_party_address }}}} (hereinafter referred to as the \"Receiving Party\", which expression shall, unless repugnant to the context or meaning thereof, be deemed to include its successors and permitted assigns)""")

doc.add_paragraph("""The Disclosing Party and the Receiving Party are collectively referred to as the \"Parties\" and individually as a \"Party\".""")

doc.add_paragraph()

# WHEREAS clauses
whereas = doc.add_paragraph()
whereas_run = whereas.add_run("WHEREAS:")
whereas_run.bold = True

doc.add_paragraph(f"A. The Parties wish to explore {{{{ purpose }}}}.")

doc.add_paragraph("B. In connection with the aforementioned purpose, the Disclosing Party may disclose certain confidential and proprietary information to the Receiving Party.")

doc.add_paragraph("C. The Parties desire to protect the confidentiality of such information and prevent its unauthorized use or disclosure.")

doc.add_paragraph()

# NOW THEREFORE
now_therefore = doc.add_paragraph()
now_therefore_run = now_therefore.add_run("NOW, THEREFORE, in consideration of the mutual covenants and agreements contained herein, and for other good and valuable consideration, the receipt and sufficiency of which are hereby acknowledged, the Parties agree as follows:")
now_therefore_run.bold = True

doc.add_paragraph()

# Clauses
def add_clause(doc, number, title, content):
    clause_heading = doc.add_paragraph()
    clause_heading_run = clause_heading.add_run(f"{number}. {title}")
    clause_heading_run.bold = True
    clause_heading_run.font.size = Pt(12)
    
    doc.add_paragraph(content)
    doc.add_paragraph()

add_clause(doc, "1", "DEFINITION OF CONFIDENTIAL INFORMATION",
"""\"Confidential Information\" means any and all information, whether written, oral, electronic, or visual, disclosed by the Disclosing Party to the Receiving Party, including but not limited to:
a) Technical data, trade secrets, know-how, research, product plans, products, services, customers, customer lists, markets, software, developments, inventions, processes, formulas, technology, designs, drawings, engineering, hardware configuration information, marketing, finances, or other business information;
b) Information that is marked as \"Confidential\", \"Proprietary\", or with a similar designation;
c) Information that would reasonably be considered confidential given the nature of the information and the circumstances of disclosure.""")

add_clause(doc, "2", "EXCLUSIONS FROM CONFIDENTIAL INFORMATION",
"""Confidential Information shall not include information that:
a) Is or becomes publicly available through no breach of this Agreement by the Receiving Party;
b) Is rightfully received by the Receiving Party from a third party without breach of any confidentiality obligation;
c) Is independently developed by the Receiving Party without use of or reference to the Confidential Information;
d) Is required to be disclosed by law, regulation, or court order, provided that the Receiving Party provides prompt notice to the Disclosing Party of such requirement.""")

add_clause(doc, "3", "OBLIGATIONS OF RECEIVING PARTY",
"""The Receiving Party agrees to:
a) Hold and maintain the Confidential Information in strict confidence;
b) Not disclose the Confidential Information to any third parties without the prior written consent of the Disclosing Party;
c) Use the Confidential Information solely for the purpose specified in this Agreement;
d) Protect the Confidential Information using the same degree of care that it uses to protect its own confidential information, but in no event less than reasonable care;
e) Limit access to the Confidential Information to its employees, contractors, and advisors who have a legitimate need to know and who have been informed of the confidential nature of such information.""")

add_clause(doc, "4", "OWNERSHIP",
"""All Confidential Information remains the sole property of the Disclosing Party. No license or right under any patent, copyright, trademark, or other intellectual property right is granted to or conferred upon the Receiving Party by this Agreement.""")

add_clause(doc, "5", "TERM AND TERMINATION",
"""This Agreement shall commence on the Effective Date and continue for a period of {{ term_years }} year(s) from the date of this Agreement, unless earlier terminated by either Party upon thirty (30) days' prior written notice to the other Party. The obligations of confidentiality shall survive termination of this Agreement for a period of two (2) years.""")

add_clause(doc, "6", "RETURN OF MATERIALS",
"""Upon termination of this Agreement or upon request by the Disclosing Party, the Receiving Party shall promptly return or destroy all Confidential Information, including all copies, notes, and derivatives thereof, and certify in writing to the Disclosing Party that such return or destruction has been completed.""")

add_clause(doc, "7", "NO WARRANTY",
"""All Confidential Information is provided \"AS IS\". The Disclosing Party makes no warranties, express, implied, or otherwise, regarding its accuracy, completeness, or performance.""")

add_clause(doc, "8", "REMEDIES",
"""The Receiving Party acknowledges that the Confidential Information is valuable and unique, and that disclosure of such information would cause irreparable harm to the Disclosing Party. Therefore, the Disclosing Party shall be entitled to seek equitable relief, including injunction and specific performance, in the event of any breach of this Agreement, in addition to all other remedies available at law or in equity.""")

add_clause(doc, "9", "GOVERNING LAW AND JURISDICTION",
f"""This Agreement shall be governed by and construed in accordance with the laws of India. The courts of {{{{ jurisdiction }}}} shall have exclusive jurisdiction over any disputes arising under this Agreement.""")

add_clause(doc, "10", "ENTIRE AGREEMENT",
"""This Agreement constitutes the entire agreement between the Parties with respect to the subject matter hereof and supersedes all prior or contemporaneous oral or written agreements concerning such subject matter.""")

add_clause(doc, "11", "AMENDMENT",
"""No amendment, modification, or waiver of any provision of this Agreement shall be effective unless in writing and signed by both Parties.""")

# Signature block
doc.add_paragraph()
doc.add_paragraph("IN WITNESS WHEREOF, the Parties have executed this Non-Disclosure Agreement as of the date first written above.")

doc.add_paragraph()
doc.add_paragraph()

# Disclosing Party signature
disc_party = doc.add_paragraph()
disc_party_run = disc_party.add_run("FOR THE DISCLOSING PARTY:")
disc_party_run.bold = True

doc.add_paragraph()
doc.add_paragraph(f"Signature: _________________________")
doc.add_paragraph(f"Name: {{{{ disclosing_party_signatory }}}}")
doc.add_paragraph(f"Title: _________________________")
doc.add_paragraph(f"Date: _________________________")

doc.add_paragraph()
doc.add_paragraph()

# Receiving Party signature
rec_party = doc.add_paragraph()
rec_party_run = rec_party.add_run("FOR THE RECEIVING PARTY:")
rec_party_run.bold = True

doc.add_paragraph()
doc.add_paragraph(f"Signature: _________________________")
doc.add_paragraph(f"Name: {{{{ receiving_party_signatory }}}}")
doc.add_paragraph(f"Title: _________________________")
doc.add_paragraph(f"Date: _________________________")

# Save
output_path = TEMPLATES_DIR / "nda.docx"
doc.save(output_path)

print(f"✅ Created NDA Jinja2 template: {output_path}")
print("\nTemplate variables:")
print("  - {{ disclosing_party }}")
print("  - {{ disclosing_party_address }}")
print("  - {{ receiving_party }}")
print("  - {{ receiving_party_address }}")
print("  - {{ effective_date }}")
print("  - {{ purpose }}")
print("  - {{ jurisdiction }}")
print("  - {{ term_years }}")
print("  - {{ disclosing_party_signatory }}")
print("  - {{ receiving_party_signatory }}")
