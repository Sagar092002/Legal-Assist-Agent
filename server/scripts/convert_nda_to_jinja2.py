"""
Convert existing NDA template to Jinja2 format
"""

from docx import Document
from pathlib import Path
import re

BASE_DIR = Path(__file__).parent.parent
TEMPLATES_DIR = BASE_DIR / "data" / "templates"

# Load the existing NDA template
input_file = TEMPLATES_DIR / "Non Disclosure Agreement.docx"
output_file = TEMPLATES_DIR / "NDA-Jinja2.docx"

print("="*80)
print("CONVERTING NDA TO JINJA2 FORMAT")
print("="*80)
print(f"\nInput: {input_file}")
print(f"Output: {output_file}\n")

doc = Document(input_file)

# Define placeholder mappings
replacements = {
    # Exact phrase replacements
    "<Party 1>": "{{ party_1_name }}",
    "<<address>>": "{{ party_1_address }}",
    "[Please fill in Customers name]": "{{ party_2_name }}",
    "[Please fill in address]": "{{ party_2_address }}",
    "[Please fill in details of proposed transaction]": "{{ proposed_transaction_details }}",
    
    # Underscores for dates (in order of appearance)
    "the _________ day of _________": "the {{ day }} day of {{ month }}, {{ year }}"
}

# Additional patterns to handle
underscore_pattern = r'_{5,}'
party_name_pattern = r'"____"'

# Track conversions
conversions = []

# Process all paragraphs
for para in doc.paragraphs:
    original_text = para.text
    new_text = original_text
    
    # Apply exact replacements
    for old, new in replacements.items():
        if old in new_text:
            new_text = new_text.replace(old, new)
            conversions.append(f"{old} → {new}")
    
    # Handle special party name placeholder
    if '"____"' in new_text:
        new_text = new_text.replace('"____"', '"{{ party_1_short_name }}"')
        conversions.append('"____" → "{{ party_1_short_name }}"')
    
    # Update paragraph text if changed
    if new_text != original_text:
        para.text = new_text

# Process tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                original_text = para.text
                new_text = original_text
                
                # Apply exact replacements
                for old, new in replacements.items():
                    if old in new_text:
                        new_text = new_text.replace(old, new)
                        conversions.append(f"{old} → {new} (in table)")
                
                # Handle special party name placeholder
                if '"____"' in new_text:
                    new_text = new_text.replace('"____"', '"{{ party_1_short_name }}"')
                    conversions.append('"____" → "{{ party_1_short_name }}" (in table)')
                
                # Update paragraph text if changed
                if new_text != original_text:
                    para.text = new_text

# Save converted document
doc.save(output_file)

print("✅ Conversions applied:")
for conv in set(conversions):
    print(f"  ✓ {conv}")

print(f"\n✅ Saved Jinja2 template: {output_file}")

print("\n" + "="*80)
print("JINJA2 VARIABLES IN NDA TEMPLATE")
print("="*80)
print("""
Template variables:
  1. {{ day }} - Day of agreement
  2. {{ month }} - Month of agreement
  3. {{ year }} - Year of agreement
  4. {{ party_1_name }} - First party full name
  5. {{ party_1_address }} - First party address
  6. {{ party_1_short_name }} - First party short reference name
  7. {{ party_2_name }} - Second party (customer) full name
  8. {{ party_2_address }} - Second party address
  9. {{ proposed_transaction_details }} - Details of proposed transaction

Next: Run update_nda_config.py to add this to template_config.json
""")
