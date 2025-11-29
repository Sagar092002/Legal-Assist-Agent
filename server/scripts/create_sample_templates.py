"""
Script to create sample .docx templates for testing template-based assembly
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_nda_template():
    """Create a Non-Disclosure Agreement template"""
    doc = Document()
    
    # Title
    title = doc.add_heading('NON-DISCLOSURE AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Intro paragraph
    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run('This Non-Disclosure Agreement ("Agreement") is entered into on ').font.size = Pt(11)
    intro.add_run('{{AGREEMENT_DATE}}').bold = True
    intro.add_run(' (the "Effective Date"), by and between:').font.size = Pt(11)
    
    # Parties
    doc.add_paragraph()
    p1 = doc.add_paragraph()
    p1.add_run('1. ').bold = True
    p1.add_run('{{PARTY_NAME_1}}').bold = True
    p1.add_run(' (hereinafter referred to as the "Disclosing Party"), and')
    
    p2 = doc.add_paragraph()
    p2.add_run('2. ').bold = True
    p2.add_run('{{PARTY_NAME_2}}').bold = True
    p2.add_run(' (hereinafter referred to as the "Receiving Party")')
    
    # Recitals
    doc.add_paragraph()
    doc.add_heading('RECITALS', level=1)
    doc.add_paragraph(
        'WHEREAS, the parties wish to explore a business opportunity of mutual interest '
        'and in connection with this opportunity, the Disclosing Party may disclose to the '
        'Receiving Party certain confidential technical and business information that the '
        'Disclosing Party desires the Receiving Party to treat as confidential.'
    )
    
    # Agreement section
    doc.add_paragraph()
    doc.add_heading('AGREEMENT', level=1)
    
    # Section 1: Purpose
    doc.add_heading('1. Purpose', level=2)
    purpose_para = doc.add_paragraph('The purpose of this Agreement is to protect confidential information related to ')
    purpose_para.add_run('{{PURPOSE}}').bold = True
    purpose_para.add_run('.')
    
    # Section 2: Definition of Confidential Information
    doc.add_heading('2. Definition of Confidential Information', level=2)
    doc.add_paragraph(
        '"Confidential Information" means all information disclosed by the Disclosing Party '
        'to the Receiving Party, whether orally, in writing, or in any other form, that is '
        'designated as confidential or that reasonably should be understood to be confidential '
        'given the nature of the information and the circumstances of disclosure.'
    )
    
    # Section 3: Obligations
    doc.add_heading('3. Obligations of Receiving Party', level=2)
    doc.add_paragraph(
        'The Receiving Party agrees to:\n'
        '(a) Hold and maintain the Confidential Information in strict confidence;\n'
        '(b) Not disclose the Confidential Information to any third parties without prior written consent;\n'
        '(c) Use the Confidential Information solely for the Purpose stated above;\n'
        '(d) Protect the Confidential Information using the same degree of care it uses to protect its own confidential information.'
    )
    
    # Section 4: Term
    doc.add_heading('4. Term', level=2)
    term_para = doc.add_paragraph('This Agreement shall remain in effect for a period of ')
    term_para.add_run('{{TERM_DURATION}}').bold = True
    term_para.add_run(' from the Effective Date, unless terminated earlier by mutual written agreement.')
    
    # Section 5: Governing Law
    doc.add_heading('5. Governing Law', level=2)
    law_para = doc.add_paragraph('This Agreement shall be governed by and construed in accordance with the laws of ')
    law_para.add_run('{{LOCATION}}').bold = True
    law_para.add_run(', without regard to its conflict of law provisions.')
    
    # Signature section
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_heading('IN WITNESS WHEREOF', level=1)
    doc.add_paragraph('The parties have executed this Agreement as of the date first above written.')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature blocks
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    # Disclosing Party
    table.cell(0, 0).text = 'DISCLOSING PARTY'
    table.cell(1, 0).text = 'Name: {{PARTY_NAME_1}}'
    table.cell(2, 0).text = 'Signature: _____________________'
    table.cell(3, 0).text = 'Date: _____________________'
    
    # Receiving Party
    table.cell(0, 1).text = 'RECEIVING PARTY'
    table.cell(1, 1).text = 'Name: {{PARTY_NAME_2}}'
    table.cell(2, 1).text = 'Signature: _____________________'
    table.cell(3, 1).text = 'Date: _____________________'
    
    return doc

def create_employment_agreement_template():
    """Create an Employment Agreement template"""
    doc = Document()
    
    # Title
    title = doc.add_heading('EMPLOYMENT AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Intro
    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run('This Employment Agreement ("Agreement") is made and entered into as of ').font.size = Pt(11)
    intro.add_run('{{START_DATE}}').bold = True
    intro.add_run(' by and between:').font.size = Pt(11)
    
    # Parties
    doc.add_paragraph()
    employer = doc.add_paragraph()
    employer.add_run('EMPLOYER: ').bold = True
    employer.add_run('{{COMPANY_NAME}}').bold = True
    employer.add_run(', a company registered under the laws of ')
    employer.add_run('{{LOCATION}}').bold = True
    
    doc.add_paragraph()
    employee = doc.add_paragraph()
    employee.add_run('EMPLOYEE: ').bold = True
    employee.add_run('{{EMPLOYEE_NAME}}').bold = True
    employee.add_run(', residing at ')
    employee.add_run('{{EMPLOYEE_ADDRESS}}').bold = True
    
    # Recitals
    doc.add_paragraph()
    doc.add_heading('RECITALS', level=1)
    doc.add_paragraph(
        'WHEREAS, the Employer desires to employ the Employee, and the Employee desires '
        'to be employed by the Employer, on the terms and conditions set forth in this Agreement.'
    )
    
    # Terms
    doc.add_paragraph()
    doc.add_heading('TERMS AND CONDITIONS', level=1)
    
    # Section 1: Position
    doc.add_heading('1. Position and Duties', level=2)
    position_para = doc.add_paragraph('The Employee is hired for the position of ')
    position_para.add_run('{{JOB_TITLE}}').bold = True
    position_para.add_run(' and shall perform the following duties: ')
    position_para.add_run('{{JOB_DUTIES}}').bold = True
    position_para.add_run('.')
    
    # Section 2: Compensation
    doc.add_heading('2. Compensation', level=2)
    salary_para = doc.add_paragraph('The Employer shall pay the Employee a salary of ')
    salary_para.add_run('{{SALARY_AMOUNT}}').bold = True
    salary_para.add_run(' per ')
    salary_para.add_run('{{SALARY_PERIOD}}').bold = True
    salary_para.add_run(', payable in accordance with the Employer\'s standard payroll practices.')
    
    # Section 3: Benefits
    doc.add_heading('3. Benefits', level=2)
    benefits_para = doc.add_paragraph('The Employee shall be entitled to the following benefits: ')
    benefits_para.add_run('{{BENEFITS}}').bold = True
    benefits_para.add_run('.')
    
    # Section 4: Working Hours
    doc.add_heading('4. Working Hours', level=2)
    hours_para = doc.add_paragraph('The Employee\'s regular working hours shall be ')
    hours_para.add_run('{{WORKING_HOURS}}').bold = True
    hours_para.add_run(' per week.')
    
    # Section 5: Leave
    doc.add_heading('5. Leave Entitlement', level=2)
    leave_para = doc.add_paragraph('The Employee shall be entitled to ')
    leave_para.add_run('{{ANNUAL_LEAVE}}').bold = True
    leave_para.add_run(' days of paid annual leave per year.')
    
    # Section 6: Term
    doc.add_heading('6. Term of Employment', level=2)
    term_para = doc.add_paragraph('This employment shall commence on ')
    term_para.add_run('{{START_DATE}}').bold = True
    term_para.add_run(' and shall continue until terminated by either party in accordance with Section 7.')
    
    # Section 7: Termination
    doc.add_heading('7. Termination', level=2)
    notice_para = doc.add_paragraph('Either party may terminate this Agreement by providing ')
    notice_para.add_run('{{NOTICE_PERIOD}}').bold = True
    notice_para.add_run(' days written notice to the other party.')
    
    # Section 8: Confidentiality
    doc.add_heading('8. Confidentiality', level=2)
    doc.add_paragraph(
        'The Employee agrees to maintain strict confidentiality regarding all proprietary '
        'information, trade secrets, and confidential data of the Employer, both during and '
        'after the term of employment.'
    )
    
    # Section 9: Governing Law
    doc.add_heading('9. Governing Law', level=2)
    law_para = doc.add_paragraph('This Agreement shall be governed by the laws of ')
    law_para.add_run('{{LOCATION}}').bold = True
    law_para.add_run('.')
    
    # Signatures
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_heading('SIGNATURES', level=1)
    
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    # Employer
    table.cell(0, 0).text = 'EMPLOYER'
    table.cell(1, 0).text = 'Company: {{COMPANY_NAME}}'
    table.cell(2, 0).text = 'Signature: _____________________'
    table.cell(3, 0).text = 'Date: _____________________'
    
    # Employee
    table.cell(0, 1).text = 'EMPLOYEE'
    table.cell(1, 1).text = 'Name: {{EMPLOYEE_NAME}}'
    table.cell(2, 1).text = 'Signature: _____________________'
    table.cell(3, 1).text = 'Date: _____________________'
    
    return doc

def create_lease_agreement_template():
    """Create a Lease Agreement template"""
    doc = Document()
    
    # Title
    title = doc.add_heading('LEASE AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Intro
    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run('This Lease Agreement ("Agreement") is made on ').font.size = Pt(11)
    intro.add_run('{{LEASE_START_DATE}}').bold = True
    intro.add_run(' between:').font.size = Pt(11)
    
    # Parties
    doc.add_paragraph()
    landlord = doc.add_paragraph()
    landlord.add_run('LANDLORD: ').bold = True
    landlord.add_run('{{LANDLORD_NAME}}').bold = True
    
    tenant = doc.add_paragraph()
    tenant.add_run('TENANT: ').bold = True
    tenant.add_run('{{TENANT_NAME}}').bold = True
    
    # Property description
    doc.add_paragraph()
    doc.add_heading('1. Property Description', level=2)
    property_para = doc.add_paragraph('The Landlord agrees to lease to the Tenant the property located at: ')
    property_para.add_run('{{PROPERTY_ADDRESS}}').bold = True
    property_para.add_run(' (the "Property").')
    
    # Lease term
    doc.add_heading('2. Lease Term', level=2)
    term_para = doc.add_paragraph('The lease term shall be for ')
    term_para.add_run('{{LEASE_DURATION}}').bold = True
    term_para.add_run(' months, commencing on ')
    term_para.add_run('{{LEASE_START_DATE}}').bold = True
    term_para.add_run(' and ending on ')
    term_para.add_run('{{LEASE_END_DATE}}').bold = True
    term_para.add_run('.')
    
    # Rent
    doc.add_heading('3. Rent', level=2)
    rent_para = doc.add_paragraph('The monthly rent shall be ')
    rent_para.add_run('{{MONTHLY_RENT}}').bold = True
    rent_para.add_run(', payable on or before the ')
    rent_para.add_run('{{RENT_DUE_DATE}}').bold = True
    rent_para.add_run(' of each month.')
    
    # Security deposit
    doc.add_heading('4. Security Deposit', level=2)
    deposit_para = doc.add_paragraph('The Tenant shall pay a security deposit of ')
    deposit_para.add_run('{{SECURITY_DEPOSIT}}').bold = True
    deposit_para.add_run(' upon execution of this Agreement.')
    
    # Utilities
    doc.add_heading('5. Utilities', level=2)
    utilities_para = doc.add_paragraph('The Tenant shall be responsible for: ')
    utilities_para.add_run('{{TENANT_UTILITIES}}').bold = True
    utilities_para.add_run('.')
    
    # Signatures
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_heading('SIGNATURES', level=1)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    table.cell(0, 0).text = 'LANDLORD'
    table.cell(1, 0).text = 'Name: {{LANDLORD_NAME}}'
    table.cell(2, 0).text = 'Signature: _____________________'
    table.cell(3, 0).text = 'Date: _____________________'
    
    table.cell(0, 1).text = 'TENANT'
    table.cell(1, 1).text = 'Name: {{TENANT_NAME}}'
    table.cell(2, 1).text = 'Signature: _____________________'
    table.cell(3, 1).text = 'Date: _____________________'
    
    return doc

def main():
    """Create all sample templates"""
    # Create directory structure
    base_dir = Path(__file__).parent.parent / 'data' / 'templates'
    
    employment_dir = base_dir / 'employment'
    property_dir = base_dir / 'property'
    
    employment_dir.mkdir(parents=True, exist_ok=True)
    property_dir.mkdir(parents=True, exist_ok=True)
    
    # Create templates
    print("📝 Creating sample templates...")
    
    # NDA Template
    nda = create_nda_template()
    nda_path = employment_dir / 'nda.docx'
    nda.save(nda_path)
    print(f"✅ Created: {nda_path}")
    
    # Employment Agreement Template
    employment = create_employment_agreement_template()
    employment_path = employment_dir / 'employment_agreement.docx'
    employment.save(employment_path)
    print(f"✅ Created: {employment_path}")
    
    # Lease Agreement Template
    lease = create_lease_agreement_template()
    lease_path = property_dir / 'lease_agreement.docx'
    lease.save(lease_path)
    print(f"✅ Created: {lease_path}")
    
    print("\n🎉 All templates created successfully!")
    print(f"\n📁 Template directory: {base_dir}")

if __name__ == '__main__':
    main()
