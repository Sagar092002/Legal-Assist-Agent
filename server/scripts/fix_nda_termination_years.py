"""
Fix NDA template - replace dots with Jinja2 variables for termination years
"""

from docx import Document
from pathlib import Path
import re

BASE_DIR = Path(__file__).parent.parent
TEMPLATES_DIR = BASE_DIR / "data" / "templates"
INPUT_FILE = TEMPLATES_DIR / "NDA-Jinja2-Fixed.docx"
OUTPUT_FILE = TEMPLATES_DIR / "NDA-Jinja2-Fixed.docx"

print("="*80)
print("FIXING NDA TERMINATION YEARS")
print("="*80)

doc = Document(INPUT_FILE)

replacements_made = 0

# Fix in paragraphs
for para in doc.paragraphs:
    if "…………….. years after the termination" in para.text:
        para.text = para.text.replace(
            "…………….. years after the termination of the binding agreement",
            "{{ termination_years }} years after the termination of the binding agreement"
        )
        replacements_made += 1
        print(f"✅ Fixed: termination years")
    
    if "……………years after the expiry" in para.text:
        para.text = para.text.replace(
            "……………years after the expiry of the binding agreement",
            "{{ expiry_years }} years after the expiry of the binding agreement"
        )
        replacements_made += 1
        print(f"✅ Fixed: expiry years")

# Fix in tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if "…………….. years after the termination" in cell.text:
                for para in cell.paragraphs:
                    if "…………….. years after the termination" in para.text:
                        para.text = para.text.replace(
                            "…………….. years after the termination of the binding agreement",
                            "{{ termination_years }} years after the termination of the binding agreement"
                        )
                        replacements_made += 1
                        print(f"✅ Fixed: termination years (in table)")
            
            if "……………years after the expiry" in cell.text:
                for para in cell.paragraphs:
                    if "……………years after the expiry" in para.text:
                        para.text = para.text.replace(
                            "……………years after the expiry of the binding agreement",
                            "{{ expiry_years }} years after the expiry of the binding agreement"
                        )
                        replacements_made += 1
                        print(f"✅ Fixed: expiry years (in table)")

# Save
doc.save(OUTPUT_FILE)

print(f"\n✅ Saved to: {OUTPUT_FILE}")
print(f"Total replacements: {replacements_made}")

# Verify
print("\n" + "="*80)
print("VERIFICATION")
print("="*80)

doc = Document(OUTPUT_FILE)
all_vars = set()

for para in doc.paragraphs:
    vars_found = re.findall(r'\{\{\s*(\w+)\s*\}\}', para.text)
    all_vars.update(vars_found)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            vars_found = re.findall(r'\{\{\s*(\w+)\s*\}\}', cell.text)
            all_vars.update(vars_found)

print(f"\n✅ All Jinja2 variables found: {sorted(all_vars)}")
print(f"Total variables: {len(all_vars)}")

# Check for remaining dots
has_dots = False
for para in doc.paragraphs:
    if "……………" in para.text:
        print(f"⚠️  Still has dots: {para.text[:100]}")
        has_dots = True

if not has_dots:
    print("\n✅ No remaining dots placeholders found!")
